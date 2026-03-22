"""Генерация DOCX-заключений.

Здесь находится построение итогового Word-документа по данным исследования:
подготовка структуры документа, секций, таблиц и итогового пути сохранения.
Если проблема проявляется уже на этапе создания `.docx`, разбирать нужно этот
модуль вместе с сервисами подготовки данных заключения.
"""

from __future__ import annotations

import re
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# --- Подготовка пути сохранения итогового DOCX ---


def _resolve_conclusion_output_path(
    output_path: str | Path | None,
    base_filename: str,
) -> Path:
    """
    Правила интерпретации output_path:

    - None -> сохранить в текущую рабочую папку с auto-именем;
    - существующая папка -> сохранить в неё с auto-именем;
    - несуществующий путь БЕЗ суффикса -> считать папкой, создать её и
      сохранить в неё с auto-именем;
    - путь С суффиксом -> считать путём к файлу;
      если суффикс не .docx, заменить его на .docx.
    """
    if output_path is None:
        return Path.cwd() / base_filename

    out = Path(output_path)

    # Существующая папка -> сохраняем внутрь папки
    if out.exists() and out.is_dir():
        return out / base_filename

    # Несуществующий путь без расширения считаем папкой
    if out.suffix == "":
        out.mkdir(parents=True, exist_ok=True)
        return out / base_filename

    # Иначе считаем, что это путь к файлу
    if out.suffix.lower() != ".docx":
        out = out.with_suffix(".docx")

    out.parent.mkdir(parents=True, exist_ok=True)
    return out


# --- Публичная генерация заключения в формате DOCX ---


def create_hla_conclusion_docx(
    class1: dict | None,
    class2: dict | None,
    num_register: int,
    head_of: str,
    acting: bool,
    biologist1: str,
    doctor: bool,
    biologist2: str,
    last_name: str,
    first_name: str,
    middle_name: str,
    screening: bool,
    clinic_name: str,
    *,
    output_path: str | Path | None = None,
    overwrite: bool = False,
) -> str:
    """
    Генерирует .docx заключение.

    Правила для output_path:
      - None:
          сохранить в текущую рабочую папку с авто-именем;
      - существующая папка:
          сохранить в неё с авто-именем;
      - несуществующий путь без расширения:
          считать папкой, создать её и сохранить в неё с авто-именем;
      - путь с расширением:
          считать полным путём к файлу.
          Если расширение не .docx, оно будет заменено на .docx.

    Возвращает путь к сохранённому файлу (строкой).
    """

    MONTHS = {
        1: "января",
        2: "февраля",
        3: "марта",
        4: "апреля",
        5: "мая",
        6: "июня",
        7: "июля",
        8: "августа",
        9: "сентября",
        10: "октября",
        11: "ноября",
        12: "декабря",
    }

    def add_paragraph(doc: Document, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        fmt = p.paragraph_format
        fmt.line_spacing = 1
        fmt.space_before = 0
        fmt.space_after = 0
        return p

    def empty_paragraph(doc: Document):
        add_paragraph(doc)

    def add_run(p, text: str, bold=False, underline=False, italic=False):
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)
        run.bold = bold
        run.underline = underline
        run.italic = italic
        return run

    def sort_specificity(antibodies: list[dict]) -> list[str]:
        if not antibodies:
            return []

        allele_groups: dict[str, set[int]] = defaultdict(set)

        for ab in antibodies:
            spec = (ab.get("Specificity") or "").strip()
            if "*" not in spec:
                continue
            gene, rest = spec.split("*", 1)
            gene = gene.strip()
            if not gene:
                continue
            group = rest.split(":", 1)[0].strip()
            if group.isdigit():
                allele_groups[gene].add(int(group))

        lines: list[str] = []
        for gene in sorted(allele_groups):
            groups = ", ".join(f"{g:02d}" for g in sorted(allele_groups[gene]))
            lines.append(f"{gene}* {groups}")

        result: list[str] = []
        for i, line in enumerate(lines):
            result.append(line + ("." if i == len(lines) - 1 else ";"))
        return result

    def add_class_block(doc: Document, data: dict):
        antibodies = data.get("antibodies") or []
        pra = data.get("pra") or ""
        hla_class = data.get("hla_class") or ""

        p = add_paragraph(doc)
        add_run(p, "Антитела к антигенам HLA ")
        add_run(p, f"{hla_class} класс: ", bold=True)

        if antibodies:
            add_run(p, "выявлены", bold=True, underline=True)
        else:
            add_run(p, "не выявлены", bold=True, underline=True)

        empty_paragraph(doc)

        if not antibodies:
            return

        p = add_paragraph(doc)
        add_run(p, "Идентификация антител к HLA:", bold=True)

        empty_paragraph(doc)

        for line in sort_specificity(antibodies):
            p = add_paragraph(doc)
            add_run(p, line)

        empty_paragraph(doc)

        p = add_paragraph(doc)
        add_run(p, f"PRA: {pra} %", bold=True)

        empty_paragraph(doc)

    def methods_line(screening_flag: bool, c1: dict | None, c2: dict | None) -> str:
        classes = []
        if c1:
            classes.append("I")
        if c2:
            classes.append("II")

        if not classes:
            return "Метод исследования: скрининг."

        classes_part = ", ".join(classes)
        if screening_flag:
            return f"Метод исследования: скрининг, LSA {classes_part} класс."
        return f"Метод исследования: LSA {classes_part} класс."

    def signature_block(
        doc: Document,
        acting_flag: bool,
        head: str,
        doctor_flag: bool,
        bio1: str,
        bio2: str,
    ):
        signature_table = doc.add_table(rows=1, cols=1)

        xml = signature_table._element
        properties = xml.tblPr

        position = OxmlElement("w:tblpPr")
        position.set(qn("w:tblpY"), "0")
        position.set(qn("w:tblpYSpec"), "bottom")
        position.set(qn("w:vertAnchor"), "margin")
        properties.append(position)

        cell = signature_table.cell(0, 0)

        head_department = (
            "И. о. зав. лабораторией HLA-типирования"
            if acting_flag
            else "Зав. лабораторией HLA-типирования"
        )

        post = "Врач к.-лаб. диагностики" if doctor_flag else "Биолог"

        sections: list[tuple[str, str]] = []

        if head:
            sections.append((head_department, head))

        if bio1:
            sections.append((post, bio1))

        if bio2:
            sections.append(("Биолог", bio2))

        if not sections:
            return

        lines: list[str] = []
        for i, (title, name) in enumerate(sections):
            if i > 0:
                lines.append("")
            lines.append(title)
            lines.append(name)

        text = "\n".join(lines)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        fmt = p.paragraph_format
        fmt.line_spacing = 1
        fmt.space_before = 0
        fmt.space_after = 0

        add_run(p, text)

    # Формируем заключение
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

    patient_full_name = f"{last_name} {first_name}"
    if middle_name:
        patient_full_name += " " + middle_name

    p = add_paragraph(doc, WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "МИНИСТЕРСТВО ЗДРАВООХРАНЕНИЯ РЕСПУБЛИКИ БЕЛАРУСЬ", bold=True)

    empty_paragraph(doc)

    p = add_paragraph(doc, WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, clinic_name, bold=True)

    p = add_paragraph(doc, WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "Лаборатория HLA-типирования", bold=True)

    empty_paragraph(doc)
    empty_paragraph(doc)

    p = add_paragraph(doc, WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "Результаты определения анти-HLA-антител", bold=True)

    empty_paragraph(doc)
    empty_paragraph(doc)
    empty_paragraph(doc)

    p = add_paragraph(doc)
    now = datetime.now()
    add_run(p, f"«{now.day:02d}» {MONTHS[now.month]} {now.year} г.")

    empty_paragraph(doc)

    p = add_paragraph(doc)
    add_run(p, f"ФИО пациента: {patient_full_name}", bold=True)

    empty_paragraph(doc)

    if class1 is None and class2 is None:
        add_class_block(doc, {"hla_class": "I", "pra": "", "antibodies": []})
        add_class_block(doc, {"hla_class": "II", "pra": "", "antibodies": []})
    else:
        if class1 is not None:
            add_class_block(doc, class1)

        if class2 is not None:
            add_class_block(doc, class2)

    p = add_paragraph(doc)
    add_run(p, methods_line(screening, class1, class2), italic=True)

    signature_block(doc, acting, head_of, doctor, biologist1, biologist2)

    initials = f"{first_name[0]}."
    if middle_name:
        initials = f"{first_name[0]}.{middle_name[0]}."

    safe_last = re.sub(r'[\\/:*?"<>|]', "_", last_name.strip() or "Пациент")
    base_filename = f"{num_register}_Ab_{safe_last}_{initials}.docx"

    out = _resolve_conclusion_output_path(output_path, base_filename)

    if out.exists() and not overwrite:
        raise FileExistsError(f"Файл уже существует: {out}")

    doc.save(out)

    return str(out)
